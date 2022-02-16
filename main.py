import os
from win32com import client
import docx
import shutil
import re
import datetime

# 步骤1 改文件名称
def rename(folder,old_name,new_name):
    for file_name in os.listdir(folder):
        source = folder + file_name
        destination = source.replace(old_name,new_name)
        os.rename(source, destination)
    print('全部文件名称已经修改完成')
    print('新文件名依次是')
    res = os.listdir(folder)
    print(res)

# 步骤2 检查文件内容,把文件里面的关键词进行修改
# 转换doc为docx
def doc2docx(fn):
    word = client.Dispatch("Word.Application") # 打开word应用程序
    doc = word.Documents.Open(fn) #打开word文件
    doc.SaveAs("{}x".format(fn), 12)#另存为后缀为".docx"的文件，其中参数12或16指docx文件
    doc.Close() #关闭原来word文件
    word.Quit()
    print("doc转化为了docx")

# 全部文件夹的转换doc为docx
def doc2docxs(folder):
    for file in os.listdir(folder):
        if file.endswith( '.doc' ) and not file.startswith( '~$' ): ##注意判断条件里面要加上不读取已经打开的word文件
            doc2docx( folder + file )

# 查找doc文件的内容，进行替换
# 此函数用于批量替换合同中需要替换的信息  doc:文件  old_info和new_info：原文字和需要替换的新文字
def info_update(doc,old_info, new_info):
    # 读取段落中的所有run，找到需替换的信息进行替换
    for para in doc.paragraphs:
        for run in para.runs:
            run.text = run.text.replace(old_info, new_info) #替换信息
    # 读取表格中的所有单元格，找到需替换的信息进行替换
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.text = run.text.replace(old_info, new_info)

# 转换docx为doc
def docx2doc(fn):
    word = client.Dispatch("Word.Application")  # 打开word应用程序
    # for file in files:
    doc = word.Documents.Open(fn)  # 打开word文件
    doc.SaveAs("{}".format(fn[:-1]), 0)  # 另存为后缀为".docx"的文件，其中参数0指doc
    doc.Close()  # 关闭原来word文件
    word.Quit()

# 将文件夹所有的docx转化为doc
def docx2docs(folder):
    for file in os.listdir(folder):
        if file.endswith( '.docx' ) and not file.startswith( '~$' ):
            docx2doc( folder + file )
    print("全部转移到了新文件夹")

# 步骤3 修改文件位置
# 按年份放到不同的文件夹里，比如2020年/1月/1月6日会议纪要
def mkdir(path):
    import os
    path = path.strip()
    path = path.rstrip("\\")
    isExists = os.path.exists(path)
    if not isExists:
        os.makedirs(path)
        return True
    else:
        return False

def removefile(folder):
    for filename in os.listdir(folder):
        if filename.endswith( '.doc' ) and not filename.startswith( '~$' ):
            try:
                str1 = filename
                m = re.search( "(\d{4}-\d{1,2}-\d{1,2})", str1 )
                print( type( m ) )
                strdate = m.group( 1 )
                print( strdate )
                time = datetime.datetime.strptime( strdate, '%Y-%m-%d' ).date()
                year = str( time.year ) + "年"
                month = str( time.month ) + "月"
                print( year, month )

                try:
                    # 定义要创建的目录
                    mkpathyear = r"H:\音乐\系统\result\\"+year
                    # 调用函数
                    mkdir( mkpathyear )
                    # 定义要创建的目录
                    mkpath = mkpathyear+r"\\"+month
                    # 调用函数
                    mkdir( mkpath )
                except:
                    pass
                try:
                    shutil.move(folder+r'\\'+filename,mkpath+r'\\'+filename)
                    print(filename+'转移成功！')
                except Exception as e:
                    print('移动失败:' + e)
            except:
                pass


if __name__=="__main__":
    folder = r'H:\音乐\系统\test1\\'
    # 步骤1 改文件名称
    rename(folder, "学习元小组会议纪要","联合实验室每周例会会议纪要")
    # 把路径中的全部docx文件都转化为doc文件
    doc2docxs(folder)
    # 步骤2 检查文件内容,把文件里面的关键词进行修改
    for file in os.listdir( folder ):
        if file.endswith( '.docx' ) and not file.startswith( '~$' ):
            source = folder + file
            print( source )
            doc = docx.Document( source )
            # 替换seminar为文献阅读分享  --注意 Seminar不能被直接识别到
            info_update( doc, 'S', '' )
            info_update( doc, 's', '' )
            info_update( doc, 'eminar', '文献阅读分享' )
            info_update( doc, '中心', '实验室' )  # 替换高精尖为实验室
            doc.save( folder + "{}".format( file.split( "/" )[-1] ) )
            print( "{}替换完成".format( file ) )
    print( '检查文件内容已全部完成' )
    # 把路径中的全部docx文件都转化为doc文件
    docx2docs(folder)
    print( '正在分类整理进文件夹ing...' )
    # 步骤3 修改文件位置
    removefile( folder )
    print('整理完毕！')


